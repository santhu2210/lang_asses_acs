"""
 To compute the level of the given word document based on grammar error,
 langauge readability, lenght of the article, journal and articel type and author nationality

 This pgm is a testing pgm for RSC client
"""
import glob2
import nltk.data
from docx import Document
import pandas as pd
import spacy
import docx2txt
import pickle
from Grammar_and_Language_v1 import Grammar_and_Language
from Topic_noun_verb import Topic_noun_verb
import re
# import win32com.client as win32
# from win32com.client import constants
from metadata_xml import metadata_xml
from xml.dom import minidom
import nltk
from tqdm import tqdm
import textstat
from glob import glob
import os
import subprocess
from os.path import dirname
import string

BASE_DIR_PATH = os.path.dirname(os.path.abspath(__file__))


def writing_preprocessed_paragraphs(file):
    """
        Preprocess the word document in the path 'file'. References are removed and paras with length less than 8 are removed
    :param file: word file path
    :return: write a new preprossed word file in the same path
    """
    list_ver = []
    # total_no_of_words=0
    document = Document(file)
    flag_reference = 0

    for para in document.paragraphs:
        list_ver.append(para.text)  # this  will  gives  all  paras into  lists values
        if ((para.text.startswith('References')) | (para.text.startswith('Notes and references'))):
            flag_reference = 1  # to check for presence of  References

    # Preprocessing steps :

    if(flag_reference):
        try:
            index_of_reference = [max([i for i, item in enumerate(list_ver) if re.search('reference', item, re.IGNORECASE)])]  # Getting index of reference from  the list
            removed_reference = list_ver[:index_of_reference[0]]  # remove everything after reference
            removed_empty_list_values = list(filter(None, removed_reference))  # Remove all  empty  values from our list of sentences
            # filter(None, removed_reference)  #####Remove all  empty  values from our list of sentences
        except:
            index_of_reference = [max([i for i, item in enumerate(list_ver) if re.search('notes and references', item, re.IGNORECASE)])
                                  ]  # Getting index of notes and references from  the list
            removed_reference = list_ver[:index_of_reference[0]]  # remove everything after reference
            removed_empty_list_values = list(
                filter(None, removed_reference))  # Remove all  empty  values from our list of sentences
    else:
        removed_empty_list_values = list(filter(None, list_ver))

    # Consider only those strings  having length  more than 8
    removed_sorter_strings = []
    for i in range(removed_empty_list_values.__len__()):
        if removed_empty_list_values[i].split(' ').__len__() >= 8:
            removed_sorter_strings.append(removed_empty_list_values[i])
        else:
            pass

    # Below function  will  remove multiple spaces strings

    def tt(w):
        if '     ' in w or bool(w) == False:
            return 'space'
        else:
            return 'no space'

    removed_mutiple_spaces_strings = []
    for i in range(removed_sorter_strings.__len__()):
        if tt(removed_sorter_strings[i]) != 'space':
            removed_mutiple_spaces_strings.append(removed_sorter_strings[i])
        else:
            pass

    # Below  we are writing paragraphs into doc files
    doc = Document()
    for i in range(removed_mutiple_spaces_strings.__len__()):
        p = doc.add_paragraph(removed_mutiple_spaces_strings[i])
        doc.save(str(file)+"__"+"preprocessed"+".docx")
        # doc.save(str(file))


def save_as_docx_windows(path):
    """
        Convert doc word file to docx word file
    :param path: a *.doc word file path
    :return: a new *.docx word file saved in the same path
    """
    # In windows OS
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)


def convert_to_docx(files):
    """
    function to iterate over the list of files for conversion
    :param path: set of *.doc word file path
    :return: none
    """
    print("------- doc to docx conversion process -------")
    for file in files:
        # If windows OS call the below fn else execute the next line after the fn call
        # save_as_docx_windows(file)
        # print (file)
        subprocess.call(['soffice', '--headless', '--convert-to', 'docx', file, '--outdir', os.path.dirname(file)])
        update_file = file+'x'
        # print ('update_path -->',update_file)
        # print ("subprocess completed")
        return update_file


def abst_ext(file):
    """
     Extracts the abstract of a given document
    :param file: a word file path
    :return: abstract of the given word document
    """
    doc = Document(file)
    for para in doc.paragraphs:
        text = (para.text).split(' ')
        if (len(text) > 50):
            return para.text
        else:
            pass


def title_ext(file):
    """
         Extracts the title of a given document
        :param file: a word file path
        :return: title of the given word document
    """
    doc = Document(file)
    title = (doc.paragraphs[0]).text
    return title


def word_count(text):
    """
    Counts the number of words in a given text
    :param text: full content of a word document
    :return: no of words in the text
    """
    words = nltk.word_tokenize(text)
    puntuation_removed_words = [x for x in words if not re.fullmatch('[' + string.punctuation + ']+', x)]
    no_of_words = puntuation_removed_words.__len__()

    if (no_of_words <= 4000):
        c = 1
    elif ((no_of_words < 8000) & (no_of_words > 4000)):
        c = 2
    else:
        c = 3
    # print (c)
    return c


def get_tokens_lengths(s):
    """
    Counts the no of words in a given sentence, by considering only the alphabets and numerics
    :param s: A sentence
    :return: No of words in a given sentence
    """
    tokenize_string = str(s).split(' ')
    remove_punctuation = list(filter(None, [re.sub(r'[^A-Za-z0-9]+', '', x) for x in tokenize_string]))
    count_of_tokens = remove_punctuation.__len__()
    return count_of_tokens


def file_name_ext(file):
    """
     Extracts the file name from the given path
    :param file: file path of a word document
    :return: word document name
    """
    tmp_doc_file_path = file
    pos = tmp_doc_file_path.rfind('\\')
    file_name = tmp_doc_file_path[(pos + 1):(pos + 11)]
    return (file_name)


def Getting_Score_Result(folder_path, abs_file):
    ######
    # Preprocessing
    ######
    # path="D:\\PycharmProjects\\CE_RSC_client\\sample files for CE\\Difficult\\*\\*.doc"
    # path=folder_path+"C8EN00688N.doc"

    # print("here folder_path -->", folder_path)
    # print("fname -->", abs_file)

    path = abs_file
    if path.endswith(".doc"):
        files = glob2.glob(path)
        # files="/home/shantakumar/Projects/context_copy_edit_RAC/sample_watcher/2/C8EN00688A.docx"
        # files = glob(os.path.join("D:\\","PycharmProjects", "CE_RSC_client","sample files for CE","DIFFICULT","*","*.doc"))
        # print ("!!!!!!!!!!! Converting doc files to docx files, if any")
        path = convert_to_docx(files)

    # path = "D:\\PycharmProjects\\CE_RSC_client\\sample files for CE 1\\EASY\\C8CE00717A\\C8CE00717A.docx"
    # path = "D:\\PycharmProjects\\CE_RSC_client\\sample files for CE 1\\Difficult\\*.docx"
    # files = glob2.glob(path)
    # files = glob(os.path.join("D:\\","PycharmProjects", "CE_RSC_client","sample files for CE","DIFFICULT","*","*.docx"))

    ######
    # Preprocessing
    ######
    # path=folder_path+"C8EN00688N.docx"
    files = glob2.glob(path)
    preprocessed_flag = 0
    for file in files:
        if (file.find("preprocessed") != -1):
            # print("Preprocessing is already carried out!!")
            preprocessed_flag = 1
            break

    if(preprocessed_flag == 0):
        for file in files:
            # print(file)
            flag = file.find('preprocessed')
            if(flag == -1):  # to eliminate preprocessed file
                # print("Preprocessing ", file)
                file_name = file.split("/")[-1]
                try:
                    writing_preprocessed_paragraphs(file)
                except Exception as e:
                    print("Bad Word Document. The specific error is ", e)

                    result = [(file, file_name, 0, 0, 0, 0, 0, 0, 0, "Bad word formatting")]
                    result_df = pd.DataFrame(result, columns=['File Path', 'File Name', 'Lang/Grammar', 'Topic', 'No of words', 'Journal code', 'author nation',
                                                              'article type', 'CE', 'Obtained Level'])  # columns=['File Path','File Name', 'Lang/Grammar', 'Topic', 'No of words', 'Journal code',\
                    #      'author nation', 'article type', 'CE', 'Obtained Level'])

                    file_topic = [(file, file_name, 0, 0, "Bad word formatting")]
                    # columns=['File Path','File Name', 'Title', 'Abstract', 'obtained Level'])
                    file_topic_df = pd.DataFrame(file_topic, columns=['File Path', 'File Name', 'Title', 'Abstract', 'obtained Level'])
                    return result_df, file_topic_df

    # path="D:\\PycharmProjects\\CE_RSC_client\\sample files for CE 1\\Difficult\\*preprocessed.docx"
    # path = "D:\\PycharmProjects\\CE_RSC_client\\sample files for CE\\EASY\\C8CE00717A\\C8CE00717A.docx__preprocessed.docx"
    # files = glob2.glob(path)
    # files = glob(os.path.join("D:\\","PycharmProjects", "CE_RSC_client","sample files for CE","DIFFICULT","*","*preprocessed.docx"))

    preprocess_path = path+"__preprocessed.docx"
    files = glob2.glob(preprocess_path)

    nlp = spacy.load('en_core_web_sm')
    grmrerr_df = pd.DataFrame()
    file_topic_df = pd.DataFrame()
    result_df = pd.DataFrame()
    result = []
    file_grammar_and_lang = []
    file_topic = []
    file_grammar_and_lang_df = pd.DataFrame()

    total_file = 0

    # for computing passive sentences
    import importlib.machinery
    loader = importlib.machinery.SourceFileLoader('report', os.path.join(BASE_DIR_PATH, 'ispassive-master', 'ispassive.py'))
    handle = loader.load_module('report')
    t = handle.Tagger()

    # to iterate through and compute level of each file
    for file in tqdm(files):
        # print("result file -->", file)

        file_name = file.split("/")[-1].split("__")[0]  # file_name_ext(file)
        file_path = file
        total_file = total_file+1

        text = docx2txt.process(file)
        doc = nlp(text)
        sents = list(doc.sents)
        total_no_sent = sents.__len__()

        # a1 - Grammar and language check();

        gl = Grammar_and_Language(file)

        # computing the number of grammar errors
        tmp_grmrerr_df, no_grmr_error = gl.grammar()
        grmrerr_df = grmrerr_df.append(tmp_grmrerr_df)
        avg_grmr_error = no_grmr_error / total_no_sent

        # computing the average number of syllabel per word
        avg_syllabel_per_word = textstat.avg_syllables_per_word(text)

        # computing the parse tree height of the given doc
        avg_prse_tree_height, percnt_len_perse_less_than_3, percnt_len_perse_between_3_6, percnt_len__perse_between_7_9,\
            percnt_len_perse_greater_than_9 = gl.parse_tree(doc)

        len_sent = 0
        sent_len_lt_ten = 0
        sent_len_gt_forty = 0
        all_distances = 0
        list_all_distances = []
        passive_sent = []

        # print("Computing sentence length & dist bw subject and root verb ...")
        for i in range(sents.__len__()):
            s = str(sents[i])
            passive_sent.append(t.is_passive(s))

            len_s = get_tokens_lengths(s)
            len_sent = len_sent + (len_s)

            if (len_s < 10):
                # print('10', len_s, s)
                sent_len_lt_ten = sent_len_lt_ten + 1
            elif (len_s > 40):
                # print('40', len_s, s)
                sent_len_gt_forty = sent_len_gt_forty + 1

            '''
            Get all Distance between subject & ROOT
            '''
            dist = gl.distnce_sub_root(s)
            all_distances = all_distances + (dist)
            list_all_distances.append(dist)

        avg_len_noun_verb = all_distances / sents.__len__()

        # computing the average no of passive sentences
        no_of_passive_sent = 0
        for p in passive_sent:
            if (str(p) == ("True")):
                no_of_passive_sent = no_of_passive_sent + 1

        avg_passive_sent = no_of_passive_sent / total_no_sent

        avg_word_per_sent = len_sent / total_no_sent
        per_sent_len_lt_ten = (sent_len_lt_ten / total_no_sent) * 100
        per_sent_len_gt_forty = (sent_len_gt_forty / total_no_sent) * 100

        df3_distance = pd.DataFrame({"dist": list_all_distances})

        dist_1 = df3_distance[df3_distance["dist"] <= 1].__len__()
        dist_2 = df3_distance[df3_distance["dist"] == 2].__len__()
        dist_3 = df3_distance[df3_distance["dist"] == 3].__len__()
        dist_4_6 = df3_distance[(df3_distance["dist"] >= 4) & (df3_distance["dist"] <= 6)].__len__()
        dist_greater_than_6 = df3_distance[df3_distance["dist"] > 6].__len__()

        perc_dist_1 = round(dist_1 / len(df3_distance) * 100, 2)
        perc_dist_2 = round(dist_2 / len(df3_distance) * 100, 2)
        perc_dist_3 = round(dist_3 / len(df3_distance) * 100, 2)
        perc_dist_4_6 = round(dist_4_6 / len(df3_distance) * 100, 2)
        perc_dist_greater_than_6 = round(dist_greater_than_6 / len(df3_distance) * 100, 2)

        tmp_results = [avg_grmr_error, avg_syllabel_per_word, avg_word_per_sent,
                       avg_prse_tree_height, percnt_len_perse_less_than_3,
                       percnt_len_perse_between_3_6, percnt_len__perse_between_7_9,
                       percnt_len_perse_greater_than_9, avg_len_noun_verb, perc_dist_1, perc_dist_2, perc_dist_3,
                       perc_dist_4_6, perc_dist_greater_than_6, avg_passive_sent]

        with open(os.path.join(BASE_DIR_PATH, 'decision_tree_model', 'dt_set1_set2_set3_depth9.pkl'), 'rb') as fid:
            dt_loaded = pickle.load(fid)
        a1 = dt_loaded.predict([tmp_results])
        # print("Score based on Grammar and language check for ", file, a1[0])

        grammar_and_lang2 = [file_path, file_name, avg_grmr_error, avg_syllabel_per_word, avg_word_per_sent,
                             avg_prse_tree_height, percnt_len_perse_less_than_3,
                             percnt_len_perse_between_3_6, percnt_len__perse_between_7_9,
                             percnt_len_perse_greater_than_9, avg_len_noun_verb, perc_dist_1, perc_dist_2, perc_dist_3,
                             perc_dist_4_6, perc_dist_greater_than_6, avg_passive_sent, a1[0]]

        file_grammar_and_lang.append(grammar_and_lang2)

        # b -no.of words

        pos = file.rfind("__preprocessed.docx")
        unprocessed_file = file[0:pos]
        # print("result file -->", unprocessed_file)
        full_text = docx2txt.process(unprocessed_file)
        b = word_count(full_text)
        # print(" Score based on no.of words extraction for ", file, b)

        x = metadata_xml()

        # tmp_doc_file_path = file
        # pos = tmp_doc_file_path.rfind('\\')
        # xml_file_path = tmp_doc_file_path[0:(pos+11)]+"_metadata.xml"
        #

        try:
            xml_path = glob(os.path.join(dirname(file), '*.xml'))
            mydoc = minidom.parse(xml_path[0])

            c = x.journal_code(mydoc)
            d = x.authors_nation(mydoc)
            e = x.type_article(mydoc)
            title = x.article_title(mydoc)
            print(" Score based on metadata extraction for ", file, c, d, e)

        except:
            c = 2
            d = 2
            e = 2
            title = title_ext(file)

        # a2 - topic detection
        abstract = abst_ext(file)
        title_and_abs = title + abstract

        doc = nlp(title_and_abs)
        # score the document based on the nouns and verbs in it
        topic = Topic_noun_verb()
        a2 = topic.term_extraction(doc)

        # print(" Score based on topic detection for ", file, a2)

        topic_list = [file_path, file_name, title, abstract, a2]
        file_topic.append(topic_list)

        ce = ((0.2 * a1[0])+(0.2 * a2) + (.25 * b) + (0.1 * c) + (.15 * d) + (.1 * e))
        # ce=((0.30 * a1[0])+(0.3 * a2) + (.2 * b) + (.2 * d))

        if (ce <= 1.66):
            obtained_level = 'EASY'
        elif (ce <= 2.33):
            obtained_level = 'INTERMEDIATE'
        else:
            obtained_level = 'DIFFICULT'

        result.append((file_path, file_name, a1[0], a2, b, c, d, e, ce, obtained_level))
        # print("Overall score", file, ce, obtained_level)

        print(result)

    result_df = pd.DataFrame(result, columns=['File Path', 'File Name', 'Lang/Grammar', 'Topic', 'No of words',
                                              'Journal code', 'author nation', 'article type', 'CE', 'Obtained Level'])
    file_grammar_and_lang_df = pd.DataFrame(file_grammar_and_lang, columns=['File Path', 'File Name', 'avg_grmr_error', 'avg_syllabel_per_word', 'avg_word_per_sent',
                                                                            'avg_parse_tree_length', 'percnt_len_perse_less_than_3', 'percnt_len_perse_between_3_6',
                                                                            'percnt_len__perse_between_7_9', 'percnt_len_perse_greater_than_9',
                                                                            'avg_len_noun_verb', 'perc_dist_1', 'perc_dist_2', 'perc_dist_3',
                                                                            'perc_dist_4_6', 'perc_dist_greater_than_6', 'avg_passive_sent', 'obtained Level'])
    file_topic_df = pd.DataFrame(file_topic, columns=['File Path', 'File Name', 'Title', 'Abstract', 'obtained Level'])

    path = folder_path+"/"+"test.xlsx"
    writer = pd.ExcelWriter(path, engine='xlsxwriter')
    result_df.to_excel(writer, 'Result')
    file_topic_df.to_excel(writer, 'Result_topic')
    file_grammar_and_lang_df.to_excel(writer, 'Result_lang&grammar')
    writer.save()

    try:
        # p = os.popen('attrib +h ' + path)
        # t = p.read()
        # p.close()
        fold_path = os.listdir(folder_path+"/")

        for item in fold_path:
            if item.endswith("__preprocessed.docx"):
                os.remove(os.path.join(folder_path+"/", item))
            elif item.endswith("st.xlsx"):
                os.rename(os.path.join(folder_path+"/", item), folder_path+"/."+item)
            else:
                pass

    except:
        print("problem in delete, hidden conversion")
        print("folder path -->", folder_path)
        print("normal file path -->", path)

    return result_df, file_topic_df
